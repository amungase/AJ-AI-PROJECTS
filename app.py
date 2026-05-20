import streamlit as st
import os
from io import BytesIO
from docx import Document
from agents import run_rfp_pipeline

# Configure Page Information
st.set_page_config(page_title="AI RFP Multi-Agent System", page_icon="🎯", layout="wide")

st.title("🎯 Enterprise RFP Auto-Creation & Enhancement System")
st.subheader("Leverage specialized Multi-Agent Crews powered by Google Gemini to transform guidelines into technical bids.")

# Sidebar Configuration for Key Privacy Access
st.sidebar.header("🔑 Authentication & Context")
api_key_input = st.sidebar.text_input("Enter Gemini API Key:", type="password")

# Prepopulate inputs with sample mock context data to allow quick test reviews
st.subheader("📋 Document Ingestion Context Layer")
col1, col2 = st.columns(2)

with col1:
    st.markdown("### **Source 1: Corporate Capabilities / Asset Base**")
    default_profile = """Company Name: CloudScale Solutions Inc.
Core Competencies: Enterprise cloud migration, multi-region architecture setup, and serverless application development.
Security Compliance: SOC2 Type II Certified, GDPR Compliant.
Past Performance: Migrated a Fortune 500 logistics provider from on-premise infrastructure to AWS/GCP across 3 distinct global zones, reducing infrastructure overhead by 42%."""
    company_context = st.text_area("Internal Responding Capabilities File:", value=default_profile, height=220)

with col2:
    st.markdown("### **Source 2: Target Client RFP Framework Requirements**")
    default_rfp = """PROJECT TITLE: Request for Proposal (RFP) - Global Infrastructure Optimization
OBJECTIVE: Secure an external technical partner to modernize our distributed web presence.
CRITICAL MANDATORY GUIDELINES:
1. Vendor must demonstrate concrete data architecture strategy detailing security compliance standards.
2. Proposal must explicitly specify high-level mitigation strategies for multi-region system uptime."""
    rfp_context = st.text_area("Target Client Requirement Specs:", value=default_rfp, height=220)

# Helper function to convert markdown text string into a native .docx stream file
def convert_to_docx(markdown_text):
    doc = Document()
    doc.add_heading("Generated RFP Proposal Response", level=1)
    for paragraph in markdown_text.split("\n\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# Action Execution Trigger
if st.button("🚀 Kickoff Multi-Agent Crew Pipeline", type="primary"):
    if not api_key_input:
        st.error("❌ Please provide a valid Google Gemini API Key in the sidebar input field to invoke the models.")
    else:
        # Inject Key directly into system environmental variables safely
        os.environ["GOOGLE_API_KEY"] = api_key_input
        
        with st.spinner("🤖 Crew Status: Agents are analyzing documents and generating the proposal response..."):
            try:
                # Invoke our Multi-Agent orchestration block
                final_proposal = run_rfp_pipeline(company_context, rfp_context)
                
                st.success("✅ Proposal Strategy Successfully Completed by Agents!")
                
                # Output Render
                st.markdown("### 📄 Final Document Draft Preview:")
                st.markdown(final_proposal)
                
                # Deliverable compilation file generation step
                docx_file = convert_to_docx(str(final_proposal))
                
                st.download_button(
                    label="📥 Download Structured Word Document (.docx)",
                    data=docx_file,
                    file_name="Automated_RFP_Response.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"An unexpected exception error occurred in the execution chain: {e}")