import streamlit as st
import os
import sys
from io import BytesIO
from docx import Document
import pypdf

# Try importing the multi-agent pipeline function
try:
    from agents import run_rfp_pipeline
except ImportError:
    st.error("Error: Could not locate 'agents.py' file in the current working directory.")
    sys.exit(1)

# App layout setup
st.set_page_config(page_title="RFP Architect V2", page_icon="🏢", layout="wide")

# Secure Verification check for the key passed via the Command Line
if "GOOGLE_API_KEY" not in os.environ or not os.environ["GOOGLE_API_KEY"].strip():
    st.error("🔒 **Security Alert:** Gemini API Key not detected in system environment.")
    st.info("To boot this production V2 app, pass your key via the command line terminal using:\n\n`set GOOGLE_API_KEY=your_key_here && streamlit run app.py` (Windows)\n\nor\n\n`export GOOGLE_API_KEY='your_key_here' && streamlit run app.py` (Mac/Linux)")
    st.stop()

st.title("🏢 Enterprise RFP Auto-Creation")
st.caption("Production-grade Multi-Agent System powered by Google Gemini 2.5 Flash & CrewAI")

# Core function to handle multiple document formats seamlessly
def extract_text_from_file(uploaded_file):
    if uploaded_file is None:
        return ""
    
    file_type = uploaded_file.name.split(".")[-1].lower()
    extracted_text = ""
    
    try:
        if file_type == "txt":
            extracted_text = str(uploaded_file.read(), "utf-8")
            
        elif file_type == "docx":
            doc = Document(uploaded_file)
            extracted_text = "\n".join([para.text for para in doc.paragraphs])
            
        elif file_type == "pdf":
            pdf_reader = pypdf.PdfReader(uploaded_file)
            extracted_text = "".join([page.extract_text() for page in pdf_reader.pages])
            
        return extracted_text
    except Exception as e:
        st.error(f"Error parsing file '{uploaded_file.name}': {str(e)}")
        return ""

# Helper to pack generated markdown cleanly back into a .docx document stream
def convert_to_docx(markdown_text):
    doc = Document()
    doc.add_heading("Automated RFP Bid Response Proposal", level=1)
    for block in markdown_text.split("\n\n"):
        if block.strip():
            doc.add_paragraph(block.strip())
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# Renders Status Widget in Sidebar to verify CLI security injection
st.sidebar.success("🔒 API Key Verified via CMD Context")
st.sidebar.markdown("""
### **System Status**
- **Core Engine:** Gemini 2.5 Flash
- **Orchestrator:** CrewAI Framework
- **Architecture:** Sequential Multi-Agent Mesh
""")

# Setup visual columns for dragging and dropping files
col1, col2 = st.columns(2)

with col1:
    st.markdown("### **1. Ingest Corporate Asset Capabilities**")
    company_file = st.file_uploader("Upload Company Credentials/Case Studies:", type=["txt", "docx", "pdf"], key="company_uploader")
    
    company_text = ""
    if company_file:
        company_text = extract_text_from_file(company_file)
        if company_text:
            st.success(f"Successfully processed {len(company_text)} characters from {company_file.name}")
            with st.expander("View Ingested Context Snippet"):
                st.text(company_text[:500] + "...")

with col2:
    st.markdown("### **2. Ingest Target Client RFP Demands**")
    rfp_file = st.file_uploader("Upload Target Client RFP Framework Paper:", type=["txt", "docx", "pdf"], key="rfp_uploader")
    
    rfp_text = ""
    if rfp_file:
        rfp_text = extract_text_from_file(rfp_file)
        if rfp_text:
            st.success(f"Successfully processed {len(rfp_text)} characters from {rfp_file.name}")
            with st.expander("View Target Requirements Snippet"):
                st.text(rfp_text[:500] + "...")

st.markdown("---")

# Main execution pathway trigger
if st.button("🚀 Process Documents & Trigger Multi-Agent Crew Pipeline", type="primary", use_container_width=True):
    if not company_text or not rfp_text:
        st.warning("⚠️ Action Blocked: Please ensure you have uploaded valid corporate context and target client RFP files before proceeding.")
    else:
        # Dynamic execution feedback window layout
        with st.status("🤖 Running Multi-Agent Execution Flow...", expanded=True) as status:
            st.write("🕵️‍♂️ **Agent 1 (Auditor):** Reviewing incoming target specifications for compliance checkboxes...")
            try:
                # Triggers background multi-agent system execution
                final_output = run_rfp_pipeline(company_text, rfp_text)
                
                status.update(label="✅ Assembly Pipeline Completed Successfully!", state="complete", expanded=False)
                
                st.markdown("### 📄 Final Document Draft Preview:")
                st.markdown(final_output)
                
                # Compiling download link stream assets
                word_stream = convert_to_docx(str(final_output))
                
                st.download_button(
                    label="📥 Download Structured Word Document (.docx)",
                    data=word_stream,
                    file_name="Automated_Enterprise_Bid_Response.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e:
                status.update(label="❌ Pipeline Failed during runtime compilation.", state="error")
                st.error(f"Execution tracking failure message: {str(e)}")   